from docx import Document
from office_utils import add_bookmark, fill_bookmark

doc = Document()
doc.add_heading('插入和填充书签测试',level=0)
p1 = doc.add_paragraph()
r1 = p1.add_run('第一段第一run\n')
r2 = p1.add_run('第一段第二run\n')
p2 = doc.add_paragraph()
r3 = p2.add_run('\n\n第二段第一run')
r4 = p2.add_run('第二run')
add_bookmark(r1,'mark1',1)
add_bookmark(r2,'mark2',2)
add_bookmark(r4,'mark4',3)

fill_bookmark(doc,'mark1','Mark1Filler')
fill_bookmark(doc,'mark2','Mark2Filler')

try:
    doc.save('../output/demo.docx')
except Exception as e:
    print(e)