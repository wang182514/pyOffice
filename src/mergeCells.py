import os
import sys

# 把上级目录加入 sys.path, 这样无论从哪个目录启动都能找到 office_utils
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from office_utils import set_doc_fonts, set_cell

CN_FONT = "宋体"
TABLE_STYLE_DEFAULT = "Table Grid"

# 输出路径: 基于本文件位置计算, 不依赖启动目录
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "output")
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "检验记录模板.docx")

doc = Document()
set_doc_fonts(doc, CN_FONT)

table = doc.add_table(3, 5)
table.style = TABLE_STYLE_DEFAULT
# 合并第一列的 3 行
table.cell(0, 0).merge(table.cell(2, 0))
# 用 set_cell 写入文字 (替代手写 runs[0].bold/align, 后者无效)
set_cell(table.cell(0, 0), "检验记录", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=CN_FONT)

doc.add_paragraph('\n' * 2)
doc.save(OUTPUT_PATH)
print(f"已生成: {OUTPUT_PATH}")
