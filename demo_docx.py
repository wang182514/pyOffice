# -*- coding: utf-8 -*-
"""
python-docx 常用读写操作演示 (详细注释版)
运行: .venv/Scripts/python.exe demo_docx.py
生成: output/demo_docx.docx   (同时会读回并打印内容)

【背景知识 1: docx 文件的本质】
.docx 其实是一个 zip 压缩包, 里面装着一堆 XML 文件:
  word/document.xml  ← 正文内容(段落/表格/书签都在这里)
  word/styles.xml    ← 样式定义(标题1/列表 等格式的"预设")
  word/theme/theme1.xml ← 主题(字体配色的全局默认值)
用 unzip 解压一个 docx 就能亲眼看到。python-docx 的本质就是
"读写这些 XML" 的封装库, 遇到高级功能没有 API 时, 可以直接操作 XML。

【背景知识 2: 对象模型 —— 一切皆段落和 run】
Document (整个文档)
 ├── sections    章节: 页边距/纸张方向/页眉页脚 等页面设置
 ├── paragraphs  段落列表(只含正文顶层段落, 不含表格内的)
 │     └── runs  文字片段: 同一段里格式不同的文字拆成多个 run
 │                例: "加粗 正常" 是 1 个段落 + 2 个 run
 └── tables      表格列表(表格与段落同级, 按出现顺序)

【重要坑: 中文字体】
python-docx 默认模板的主题文件里, 中文字体是空值(<a:ea typeface=""/>)。
而 Word 渲染中文用的是 eastAsia 槽位的字体 —— 为空就找不到字形,
打开文档时中文全变成"方框叉"。解决办法: 显式设置 w:eastAsia, 见 helper。

Word 的字体其实分 4 个槽位:
  w:ascii    英文半角字符 (a,b,c,1,2,3)
  w:hAnsi    高 ANSI 字符 (带音标的欧洲文字等)
  w:eastAsia 中日韩字符   ← 中文看这个!
  w:cs       复杂文种 (阿拉伯文等)
run.font.name = "微软雅黑" 只会填 ascii + hAnsi 两个槽位,
所以中文必须再手动补 eastAsia 槽位。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor          # Pt=磅(字号) Cm=厘米(尺寸)
from docx.enum.text import WD_ALIGN_PARAGRAPH     # 段落对齐方式枚举
from docx.enum.table import WD_TABLE_ALIGNMENT    # 表格对齐方式枚举
from docx.oxml.ns import qn                       # qn(): 把 "w:id" 转成带命名空间的完整标签
# 书签操作已封装到 office_utils 包, 从包导入 (包内函数不再在本文件定义)
from office_utils import (
    set_run_font, set_style_font,
    add_bookmark, find_bookmark, fill_bookmark,
)

OUT_DIR = "output"
os.makedirs(OUT_DIR, exist_ok=True)               # 输出目录不存在就创建
DOCX_PATH = os.path.join(OUT_DIR, "demo_docx.docx")

CN_FONT = "微软雅黑"   # 中文字体, 可换成 "宋体"/"等线" 等(需本机装有该字体)


# =================================================================
# 表格样式速查 (python-docx 默认模板内置 100 种表格样式)
# =================================================================
# 命名规律: 系列 + (可选)"Accent N"(强调色编号)
#
#   "Table Grid"                 网格型: 黑色细边框、无底色 ← Word 最"默认"的表格
#   "Normal Table"               普通表格: 完全没有边框和底色
#   "Light Grid"                 浅色网格: 极浅灰细边框
#   "Light Grid Accent N"        浅色网格-强调N: 彩色边框 + 表头底纹
#   "Light List Accent N"        浅色列表-N: 只有横向细线 + 彩色表头
#   "Light Shading Accent N"     浅色底纹-N: 表头浅底色, 数据行隔行极浅底色
#   "Medium Shading 1/2 Accent N" 中度底纹-N: 表头深底色白字, 隔行浅底色
#   "Medium List 1/2 Accent N"   中度列表-N: 表头深底色, 仅横向边框
#   "Medium Grid 1/2/3 Accent N" 中度网格-N: 边框比 Light 系列粗
#   "Dark List/Shading Accent N" 深色-N: 整个表格深底色白字(慎用)
#   "Colorful Shading/List/Grid Accent N" 彩色-N: 每行颜色都不同
#
# Accent 颜色对照: 1=蓝 2=红 3=绿 4=紫 5=青 6=橙 (Office 经典配色)
# 记忆点: 名字不带 "Accent" 的(Table Grid/Normal Table/Light Grid)是黑白灰。
TABLE_STYLE_DEFAULT = "Table Grid"          # 黑色细边框、无底色
TABLE_STYLE_BLUE = "Light Grid Accent 1"    # 蓝色系底纹样式


# =================================================================
# 书签操作已封装到 office_utils 包, 见 office_utils/bookmark.py
# (add_bookmark / find_bookmark / fill_bookmark / fill_bookmarks)
# 上方的 XML 结构说明和工作流说明仍可参考。
# =================================================================


# ---------------------------------------------------------------
# 一、写入操作 (生成一个"模板"文档)
# ---------------------------------------------------------------
def write_docx(path):
    # Document() 不带参数 = 基于 python-docx 自带的默认模板新建空文档
    doc = Document()

    # 0. 关键修复: 给本文件用到的所有样式统一指定中文字体
    #    (不设置的话, 主题里中文字体为空 → 打开显示"方框叉")
    #    在样式层面设一次, 后面所有用这些样式的段落自动继承, 比
    #    每个 run 单独设省事。styles["名字"] 按样式名取样式对象。
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2",
                       "List Bullet", "List Number"]:
        set_style_font(doc.styles[style_name], CN_FONT)

    # 1. 标题: add_heading(文字, level)
    #    level=0 → Title 样式(文档大标题, 通常是更大的字)
    #    level=1~9 → "标题 1~9" 样式(用于导航大纲/自动目录)
    doc.add_heading("python-docx 读写演示", level=0)
    doc.add_heading("一、段落与文字格式", level=1)

    # 2. 普通段落: 默认 Normal 样式
    doc.add_paragraph("这是一个普通段落，默认样式。")

    # 3. 带格式的段落 —— 核心: 格式加在 run 上, 不是段落上!
    #    一段文字想出现多种格式(部分加粗/部分红字), 就拆成多个 run
    p = doc.add_paragraph()
    r1 = p.add_run("加粗 ")
    r1.bold = True                           # 加粗 (取消用 r1.bold = None)
    r2 = p.add_run("斜体 ")
    r2.italic = True                         # 斜体
    r3 = p.add_run("红色字 ")
    r3.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)   # 颜色: RGB 三元组
    r4 = p.add_run("16号字")
    r4.font.size = Pt(16)                    # 字号: Pt(磅), Word 里的"号"要换算
    # 统一设置中文字体 (含中文的 run 都要走 set_run_font, 理由见文件头)
    for r in p.runs:
        set_run_font(r, CN_FONT)

    # 4. 对齐方式: 常用 LEFT / CENTER / RIGHT / JUSTIFY(两端对齐)
    p = doc.add_paragraph("居中显示的段落")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 5. 项目符号列表: style 参数指定样式名(就是 Word 样式面板里的名字)
    doc.add_paragraph("列表项一", style="List Bullet")   # 圆点列表
    doc.add_paragraph("列表项二", style="List Bullet")

    # 6. 编号列表
    doc.add_paragraph("第一步", style="List Number")     # 1. 2. 3. 列表
    doc.add_paragraph("第二步", style="List Number")

    # 7. 表格
    doc.add_heading("二、表格", level=1)
    headers = ["姓名", "部门", "工资"]
    rows = [
        ["张三", "研发部", "15000"],
        ["李四", "市场部", "12000"],
        ["王五", "人事部", "10000"],
    ]

    # 表格 1: Table Grid —— 黑色细边框、无底色(最常用的"默认"样式)
    # add_table(rows=1, cols=3): 只建骨架(1行3列空格子), 数据后面再填
    table = doc.add_table(rows=1, cols=3)
    table.style = TABLE_STYLE_DEFAULT         # 换样式只需改这一行, 清单见顶部注释
    fill_table(table, headers, rows)

    doc.add_paragraph()                       # 空段落隔开两个表(不然会粘在一起)

    # 表格 2: 带底色的样式做对比 —— 同样的数据, 不同 style 的效果
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = TABLE_STYLE_BLUE           # 想换颜色改 Accent 2~6
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER   # 整个表格在页面上居中
    fill_table(table2, headers, rows)

    # 8. 分页符: 后面的内容从新的一页开始
    doc.add_page_break()

    # 9. 页眉 / 页脚 (想恢复演示就取消注释)
    # doc.add_heading("三、页眉页脚", level=1)
    # section = doc.sections[0]                       # 第一个章节
    # section.header.paragraphs[0].text = "这是页眉"  # 页眉也是由段落组成
    # section.footer.paragraphs[0].text = "第 1 页"

    # 10. 书签: 造两个"占位段落"并打上书签, 模拟模板里的待填位置
    #     段落结构: [前缀 run][书签包住的占位符 run]
    doc.add_heading("三、书签写入", level=1)
    p = doc.add_paragraph()
    p.add_run("报告人：")                     # 前缀(在书签外, 填充时不会被删)
    add_bookmark(p.add_run("【待填写】"), "reporter", bookmark_id=100)
    p = doc.add_paragraph()
    p.add_run("报告日期：")
    add_bookmark(p.add_run("【待填写】"), "report_date", bookmark_id=101)

    # save: 序列化成 zip 写入磁盘; 已存在会直接覆盖
    doc.save(path)
    print(f"[写入完成] {path}")


def fill_table(table, headers, rows):
    """通用填表: 第 0 行当表头, 其余逐行 append
    单元格本质是"小方格里的段落", cell.text = "xx" 会替换单元格内全部文字"""
    for i, h in enumerate(headers):           # 表头
        table.rows[0].cells[i].text = h
    for row in rows:                          # 数据: add_row() 在表尾加一行
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


# ---------------------------------------------------------------
# 二、按书签写入 (模拟真实场景: 打开模板 → 填数据 → 保存)
# ---------------------------------------------------------------
def fill_template(path):
    # Document(路径) = 打开已有文档 (Document() 才是新建)
    doc = Document(path)
    fill_bookmark(doc, "reporter", "王五", font=CN_FONT)
    fill_bookmark(doc, "report_date", "2026-08-14", font=CN_FONT)
    # 保存到新路径可以保留原始模板反复使用, 这里图省事原地覆盖:
    # doc.save(os.path.join(OUT_DIR, "demo_docx_filled.docx))
    doc.save(path)
    print("[书签填充完成] reporter=王五, report_date=2026-08-14")


# ---------------------------------------------------------------
# 三、读取操作
# ---------------------------------------------------------------
def read_docx(path):
    doc = Document(path)

    # 1. 读段落: doc.paragraphs 是正文顶层段落(表格内/页眉页脚的段落不在其中)
    #    para.text —— 段落纯文本(所有 run 的文字拼起来)
    #    para.style.name —— 样式名, 可用来筛选(比如只处理标题)
    print("\n--- 读取段落 (前 10 个非空) ---")
    count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:                              # 跳过空段落
            print(f"  [{para.style.name}] {text}")
            count += 1
            if count >= 10:
                break

    # 2. 读表格: 行列双层遍历
    #    table.style.name 可以查到表格用的样式
    print("\n--- 读取表格 ---")
    for t_idx, table in enumerate(doc.tables):
        print(f"  表格 {t_idx + 1}  样式={table.style.name}  "
              f"({len(table.rows)} 行 x {len(table.columns)} 列):")
        for row in table.rows:                # row.cells: 这一行的单元格列表
            cells = [cell.text for cell in row.cells]
            print("    " + " | ".join(cells))
    # 注意: 合并单元格会让被合并的格子返回相同内容(重复出现)

    # 3. 读书签: 复用 find_bookmark 确认填充结果
    print("\n--- 读取书签 ---")
    for name in ["reporter", "report_date"]:
        start, end = find_bookmark(doc, name)
        if start is not None:
            para_el = start.getparent()       # 书签的父节点 = 所在段落的 <w:p>
            # 遍历段落里所有 <w:t> 文本节点拼出整段文字
            text = "".join(t.text or "" for t in para_el.iter(qn("w:t")))
            print(f"  书签 '{name}' 存在, 所在段落文字: {text}")

    # 4. 其他常用统计
    print("\n--- 其他统计 ---")
    print(f"  段落总数: {len(doc.paragraphs)}")
    print(f"  表格总数: {len(doc.tables)}")
    print(f"  图片总数: {len(doc.inline_shapes)}")    # 行内图片
    print(f"  章节总数: {len(doc.sections)}")


if __name__ == "__main__":
    # 完整流程分三步, 模拟真实的"模板填充"工作流:
    write_docx(DOCX_PATH)      # 第 1 步: 生成带书签占位符的"模板"
    fill_template(DOCX_PATH)   # 第 2 步: 打开模板, 按书签填入数据
    read_docx(DOCX_PATH)       # 第 3 步: 读回来验证结果
    path = os.path.join(OUT_DIR, "my_demo.docx")
    doc = Document(path)
    fill_bookmark(doc,"mark1","mk1")
    fill_bookmark(doc, "mark2", "mk2")
    doc.save(path)
    pass


