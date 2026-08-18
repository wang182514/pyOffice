# -*- coding: utf-8 -*-
"""
表格操作模块 —— 封装 docx 表格的常见读写

设计原则:
- 模板填充场景: "表格已存在, 按坐标填值" 一类操作(demo_docx.py 之外的主场景)
- 全代码生成场景: "从零建表 + 填表头 + 追加数据行" 一类操作(沿用 fill_table)
- 所有格式化函数都返回 cell, 方便链式调用
"""
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


# =================================================================
# 从零建表
# =================================================================
def fill_table(table, headers, rows):
    """通用填表: 第 0 行当表头, 其余逐行 append

    适合"全代码生成"场景 —— 从空白表格开始填。
    模板填充场景请用 set_cell() / add_rows()。

    Args:
        table:  python-docx 的 Table 对象
        headers: 表头列表, 写入第 0 行
        rows:    数据二维列表, 每行追加到表尾

    单元格本质是"小方格里的段落", cell.text = "xx" 会替换单元格内全部文字
    """
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if val is None else str(val)


# =================================================================
# 模板填充: 定位
# =================================================================
def find_table_by_header(doc, keyword, table_index=0):
    """按表头关键字定位模板中的表格

    模板里常有多个表格, 按 doc.tables[0] / [1] 序号找容易错位
    (插入/删除表格后序号都会变)。
    按表头文本中的关键字定位更稳。

    Args:
        doc:         python-docx 的 Document 对象
        keyword:     表头中需包含的关键字, 如 "姓名" / "检验记录"
        table_index: 同一关键字出现多次时, 取第几个 (0-based)

    Returns:
        命中的 Table 对象; 找不到返回 None

    Example:
        doc = Document("template.docx")
        table = find_table_by_header(doc, "检验记录")
        if table is not None:
            set_cell(table, 1, 0, "射频组件A")
    """
    hits = 0
    for table in doc.tables:
        first_row_text = "".join(cell.text for cell in table.rows[0].cells)
        if keyword in first_row_text:
            if hits == table_index:
                return table
            hits += 1
    return None


def find_cell_by_text(table, keyword, row_index=0):
    """在指定行的单元格中, 按文本查找单元格

    适合"先找到这一行, 再按文本定位某列"的场景

    Args:
        table:     Table 对象
        keyword:   要查找的文本 (精确匹配)
        row_index: 在哪一行查找, 默认第 0 行 (表头)

    Returns:
        命中的 Cell 对象; 找不到返回 None
    """
    if row_index >= len(table.rows):
        return None
    for cell in table.rows[row_index].cells:
        if cell.text == keyword:
            return cell
    return None


# =================================================================
# 模板填充: 写入
# =================================================================
def set_cell(cell, text, font=None, bold=None, align=None, size=None):
    """设置单元格内容(替换式), 可选字体/加粗/对齐/字号

    替换单元格内全部文字(等同于 cell.text = text),
    但在替换前先清空旧 run, 避免产生多个空 run。

    Args:
        cell:   python-docx 的 Cell 对象
        text:   要写入的文字
        font:   中文字体, 如 "宋体" (会同时设西文 + 中文槽位)
        bold:   True / False / None (不修改)
        align:  WD_ALIGN_PARAGRAPH.CENTER / LEFT / RIGHT / JUSTIFY
        size:   Pt(12) 等

    Returns:
        cell 自身, 便于链式 .font

    常见坑: 单元格对齐是"段落"属性, 不是"run"属性。
    错误: cell.paragraphs[0].runs[0].align = 'right'  ← 错! run 没 align
    正确: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    """
    # 1. 清空旧内容
    cell.text = ""

    # 2. 写入新文字
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    if bold is not None:
        run.bold = bold
    if font is not None:
        from .docx_font import set_run_font
        set_run_font(run, font)
    if size is not None:
        run.font.size = size
    return cell


def format_cell(cell, bold=None, align=None, font=None, size=None):
    """格式化单元格(不修改文字)

    已有内容的单元格只想调格式(比如表头加粗居中)时用这个。

    Args:
        cell:  Cell 对象
        bold/align/font/size: 不传的参数保持不变

    Returns:
        cell 自身, 便于链式
    """
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    for run in p.runs:
        if bold is not None:
            run.bold = bold
        if font is not None:
            from .docx_font import set_run_font
            set_run_font(run, font)
        if size is not None:
            run.font.size = size
    return cell


def format_row(row, bold=None, align=None, font=None, size=None):
    """格式化整行(不修改文字)

    适合作为 add_rows 的 formatter 回调 —— formatter 收到的是 row 不是 cell,
    用这个能把整行统一格式化 (比如所有数据行居中)。

    Args:
        row:  Row 对象
        其他参数同 format_cell

    Returns:
        row 自身
    """
    for cell in row.cells:
        format_cell(cell, bold=bold, align=align, font=font, size=size)
    return row


def add_rows(table, data, start_row=None, formatter=None):
    """在表格末尾追加数据行(不动表头)

    模板填充场景: 模板表已建好, 只需要追加数据。
    每次调用都在 add_row() 末尾追加, 适合"数据量每次不同"的动态表。

    Args:
        table:      Table 对象
        data:       二维列表, 每一项是一行数据 [val1, val2, ...]
        start_row:  数据起始行号; 缺省 = 追加到当前末尾 (即 len(table.rows))
        formatter:  可选回调, 签名 fn(row) -> row, 对每行格式化
                    例: lambda row: format_cell(row.cells[0], bold=True, align=CENTER)

    Returns:
        新增的行的列表 (供后续单独调整)

    Example:
        rows = add_rows(table, [["A", 1.5], ["B", 2.3]], start_row=1,
                        formatter=lambda r: format_cell(r, align=CENTER))
    """
    if start_row is None:
        start_row = len(table.rows)

    new_rows = []
    for row_data in data:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = "" if val is None else str(val)
        new_rows.append(table.rows[start_row])
        start_row += 1

    if formatter is not None:
        for row in new_rows:
            formatter(row)
    return new_rows


# =================================================================
# 表格表头重复（跨页表格必备）
# =================================================================
def set_repeat_header(table, row=0):
    """让表格的指定行在跨页时自动重复（类似 Word 里"重复标题行"）

    python-docx 原生 API 没暴露这个功能, 需要直接操作 XML。
    常见于"数据表跨页" —— 模板里数据表跨页时, 第 2 页起缺失表头,
    加了这行后 Word 会在每页顶部自动重复表头。

    Args:
        table: Table 对象
        row:   要重复的行号 (默认 0 = 表头)
    """
    if row >= len(table.rows):
        raise ValueError(f"行号 {row} 超出范围 (表格共 {len(table.rows)} 行)")

    tr = table.rows[row]._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:                         # 没有 trPr 就创建
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    tblHeader = OxmlElement("w:tblHeader")
    # 注意: 先删旧元素再 append, 避免重复加入
    for old in trPr.findall(qn("w:tblHeader")):
        trPr.remove(old)
    trPr.append(tblHeader)
