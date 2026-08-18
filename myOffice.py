from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from office_utils import set_doc_fonts, set_run_font, set_cell, add_rows

CN_FONT = "宋体"

doc = Document()
set_doc_fonts(doc, CN_FONT)                       # 一行搞定字体初始化

doc.add_heading("大标题", level=0)
doc.add_heading("一级标题", level=1)

p = doc.add_paragraph('这是一个段落，默认样式')
p0 = p.insert_paragraph_before('插入到段落前')

p2 = doc.add_paragraph()
r1 = p2.add_run("加粗 "); r1.bold = True
r2 = p2.add_run("斜体 "); r2.italic = True
r3 = p2.add_run("红色字 "); r3.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
r4 = p2.add_run("16号字"); r4.font.size = Pt(16)
for run in p2.runs:
    set_run_font(run, CN_FONT)

pCenter = doc.add_paragraph("居中显示的段落")
pCenter.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 表格: 用新封装的 set_cell / add_rows
doc.add_heading("表格", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
set_cell(table.cell(0, 0), "姓名", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=CN_FONT)
set_cell(table.cell(0, 1), "部门", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=CN_FONT)
set_cell(table.cell(0, 2), "工资", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=CN_FONT)
add_rows(table, [
    ["张三", "研发部", 15000],
    ["李四", "市场部", 12000],
    ["王五", "人事部", 10000],
])

doc.save('demo.docx')
